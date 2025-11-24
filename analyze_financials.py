import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import unicodedata

import pandas as pd
import pdfplumber
from docx import Document

# Katalog bazowy ze sprawozdaniami
SPRAWOZDANIA_DIR = Path("pobrane/sprawozdania_2024")
SUMMARY_XLSX = Path("raporty/raport_finansowy_2024.xlsx")
ISSUES_DOCX = Path("raporty/uwagi_nieprawidlowosci.docx")
REGISTRY_FILE = Path("pobrane/Wykaz_szkół_i_placówek_oświatowych_30.09.2024_.xlsx")

POWIAT_FILTER = "raciborsk"
MIASTO_FILTER = "Racibórz"
TRANSLIT_MAP = str.maketrans(
    {
        "ą": "a",
        "ć": "c",
        "ę": "e",
        "ł": "l",
        "ń": "n",
        "ó": "o",
        "ś": "s",
        "ź": "z",
        "ż": "z",
        "Ą": "a",
        "Ć": "c",
        "Ę": "e",
        "Ł": "l",
        "Ń": "n",
        "Ó": "o",
        "Ś": "s",
        "Ź": "z",
        "Ż": "z",
    }
)

Number = Optional[float]


def clean_label(text: str) -> str:
    """Zamień wielokrotne spacje i nowe linie na pojedyncze spacje."""
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_number(text: str) -> Number:
    """Zamień zapis typu '1 234 567,89' lub '-1 234,00' na float."""
    if not text:
        return None
    text = text.replace("\xa0", " ").replace(" ", "")
    text = text.replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def extract_numbers(cells: List[str]) -> List[float]:
    """Wyciągnij wszystkie wartości liczbowe z listy komórek w kolejności wystąpienia."""
    nums: List[float] = []
    for cell in cells:
        if not cell:
            continue
        for match in re.findall(r"-?\d[\d\s\xa0,]*\d", cell):
            num = parse_number(match)
            if num is not None:
                nums.append(num)
    return nums


def parse_rzis_pdf(path: str) -> List[Dict[str, Optional[float]]]:
    """Zwróć listę wierszy: label, prev_year, current_year."""
    rows: List[Dict[str, Optional[float]]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for raw_row in table:
                    if not raw_row:
                        continue
                    label_cell = raw_row[0] or ""
                    label = clean_label(label_cell)
                    # pomijamy nagłówki bez etykiety
                    if not label:
                        continue
                    numeric_cells = [c or "" for c in raw_row[1:]]
                    numbers = extract_numbers(numeric_cells)
                    prev_val = numbers[0] if len(numbers) >= 2 else (numbers[0] if len(numbers) == 1 else None)
                    curr_val = numbers[-1] if numbers else None
                    rows.append({"label": label, "prev_year": prev_val, "current_year": curr_val})
    return rows


def find_value(rows: List[Dict[str, Optional[float]]], prefix: str) -> Tuple[Number, Number]:
    """Szukaj pierwszego wiersza, którego etykieta zaczyna się od prefix."""
    for row in rows:
        if row["label"].startswith(prefix):
            return row["prev_year"], row["current_year"]
    return None, None


def build_summary(rows: List[Dict[str, Optional[float]]]) -> Dict[str, Number]:
    """Przygotuj kluczowe agregaty kosztów/przychodów."""
    summary: Dict[str, Number] = {}
    summary["przychody_netto"] = find_value(rows, "A. Przychody netto z podstawowej działalności operacyjnej")[1]
    summary["dotacje_podstawowe"] = find_value(rows, "A.V. Dotacje na finansowanie działalności podstawowej")[1]
    summary["przychody_budzetowe"] = find_value(rows, "A.VI. Przychody z tytułu dochodów budżetowych")[1]
    summary["koszty_operacyjne"] = find_value(rows, "B. Koszty działalności operacyjnej")[1]
    summary["amortyzacja"] = find_value(rows, "B.I. Amortyzacja")[1]
    summary["materialy_i_energia"] = find_value(rows, "B.II. Zużycie materiałów i energii")[1]
    summary["uslugi_obce"] = find_value(rows, "B.III. Usługi obce")[1]
    summary["podatki_i_oplaty"] = find_value(rows, "B.IV. Podatki i opłaty")[1]
    summary["wynagrodzenia"] = find_value(rows, "B.V. Wynagrodzenia")[1]
    summary["ubezpieczenia_i_swiadczenia"] = find_value(
        rows, "B.VI. Ubezpieczenia społeczne i inne świadczenia dla pracowników"
    )[1]
    summary["pozostale_koszty_rodzajowe"] = find_value(rows, "B.VII. Pozostałe koszty rodzajowe")[1]
    summary["pozostale_przychody_operacyjne"] = find_value(rows, "D. Pozostałe przychody operacyjne")[1]
    summary["pozostale_koszty_operacyjne"] = find_value(rows, "E. Pozostałe koszty operacyjne")[1]
    summary["zysk_strata_netto"] = find_value(rows, "L. Zysk (strata) netto")[1]
    return summary


def detect_issues(name: str, summary: Dict[str, Number], student_count: Number) -> List[str]:
    """Zwróć listę uwag dla danej placówki."""
    issues: List[str] = []
    net = summary.get("zysk_strata_netto")
    costs = summary.get("koszty_operacyjne")
    revenues = summary.get("przychody_netto")

    if net is not None and net < 0:
        issues.append(f"Wynik netto ujemny ({net:,.2f} PLN).")
    if costs is not None and revenues is not None and costs > revenues:
        issues.append("Koszty operacyjne przewyższają przychody podstawowe (deficyt operacyjny).")
    if summary.get("pozostale_koszty_operacyjne"):
        issues.append("Występują pozostałe koszty operacyjne – warto sprawdzić ich naturę.")
    if summary.get("pozostale_koszty_rodzajowe"):
        issues.append("Odnotowano pozostałe koszty rodzajowe > 0.")
    if student_count is None:
        issues.append("Brak liczby uczniów/wychowanków – koszt na ucznia nie został policzony.")
    return issues


def normalize_name_from_dir(dir_path: str) -> str:
    """Zamień nazwę katalogu (slug) na czytelną nazwę."""
    base = os.path.basename(dir_path)
    name = base.replace("_", " ")
    # korekta kilku literówek wynikających z transliteracji (ł -> l)
    name = name.replace("Szkoa", "Szkola")
    name = name.replace("Zespo", "Zespol")
    name = name.replace("Zobkow", "Zlobkow")
    return name


def normalize_ascii(text: str) -> str:
    """Uprość tekst do ascii/lowercase, usuń nadmiarowe znaki."""
    if not text:
        return ""
    base = text.translate(TRANSLIT_MAP)
    normalized = unicodedata.normalize("NFKD", base)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"[^0-9A-Za-z]+", " ", ascii_text)
    return cleaned.strip().lower()


def classify_facility_type(name: str) -> str:
    norm = normalize_ascii(name)
    if norm.startswith("przedszkole"):
        return "Przedszkole"
    if norm.startswith("szkola podstawowa"):
        return "Szkola podstawowa"
    if "szkolno przedszkolny" in norm or "szkolno-przedszkolny" in norm or "szkolno rzedszkolny" in norm:
        return "Zespol szkolno-przedszkolny"
    if "zlobk" in norm or "zlob" in norm:
        return "Zlobek / Zespol Zlobkow"
    return "Inne"


def load_registry_index() -> Dict[str, Dict[str, float]]:
    """Zbuduj słownik liczby uczniów z wykazu (miasto Racibórz)."""
    base_index: Dict[str, Dict[str, float]] = {"przedszkole": {}, "szkola_podstawowa": {}, "zsp": {}, "zlobek": None}
    if not REGISTRY_FILE.exists():
        return base_index

    df = pd.read_excel(REGISTRY_FILE)
    df = df[df["Powiat"].str.contains(POWIAT_FILTER, case=False, na=False)]
    df = df[df["Gmina"].str.contains(MIASTO_FILTER, case=False, na=False)]
    df["norm_name"] = df["Nazwa placówki"].apply(normalize_ascii)

    for _, row in df.iterrows():
        ucz = row.get("ucz_ogolem")
        if pd.isna(ucz):
            continue
        try:
            count = int(float(ucz))
        except (TypeError, ValueError):
            continue
        norm = row["norm_name"]
        if m := re.search(r"przedszkole\s+nr\s*(\d+)", norm):
            base_index["przedszkole"][m.group(1)] = count
            continue
        if m := re.search(r"szkola\s+podstawowa\s+nr\s*(\d+)", norm):
            base_index["szkola_podstawowa"][m.group(1)] = count
            continue
        if m := re.search(r"zespol\s+szkolno[^\d]*nr\s*(\d+)", norm):
            base_index["zsp"][m.group(1)] = count
            continue
        if "zlob" in norm:
            base_index["zlobek"] = count
    return base_index


def match_student_count(facility_name: str, registry_index: Dict[str, Dict[str, float]]) -> Number:
    """Znajdź liczbę uczniów dla placówki na podstawie numeru/typu."""
    norm = normalize_ascii(facility_name)
    if m := re.search(r"przedszkole\s+nr\s*(\d+)", norm):
        return registry_index.get("przedszkole", {}).get(m.group(1))
    if m := re.search(r"szkola\s+podstawowa\s+nr\s*(\d+)", norm):
        return registry_index.get("szkola_podstawowa", {}).get(m.group(1))
    if m := re.search(r"zespol\s+szkolno[^\d]*nr\s*(\d+)", norm):
        return registry_index.get("zsp", {}).get(m.group(1))
    if "zlob" in norm:
        return registry_index.get("zlobek")  # type: ignore[index]
    return None


def collect_rzis_files() -> List[str]:
    """Zbierz ścieżki do RZiS 2024, z fallbackiem do pobrane/."""
    base_dirs: List[Path] = []
    if SPRAWOZDANIA_DIR.exists():
        base_dirs.append(SPRAWOZDANIA_DIR)
    pobrane_dir = Path("pobrane")
    if pobrane_dir.exists() and pobrane_dir not in base_dirs:
        base_dirs.append(pobrane_dir)

    files: List[str] = []
    seen = set()
    for base in base_dirs:
        for pdf_path in base.rglob("*.pdf"):
            lower = pdf_path.name.lower()
            if "rachunek" in lower and "2024" in lower:
                resolved = pdf_path.resolve()
                if resolved in seen:
                    continue
                seen.add(resolved)
                files.append(str(pdf_path))
    files.sort()
    if not files:
        raise SystemExit("Nie znaleziono plików Rachunek*.pdf w podkatalogach pobrane.")
    return files


def main():
    report_rows = []
    per_facility_tables: Dict[str, pd.DataFrame] = {}
    issues: Dict[str, List[str]] = {}

    registry_index = load_registry_index()
    rzis_files = collect_rzis_files()

    for pdf_path in rzis_files:
        facility_dir = os.path.dirname(pdf_path)
        facility_name = normalize_name_from_dir(facility_dir)
        facility_type = classify_facility_type(facility_name)
        student_count = match_student_count(facility_name, registry_index)
        rows = parse_rzis_pdf(pdf_path)
        summary = build_summary(rows)
        summary["liczba_uczniow"] = student_count
        summary["typ"] = facility_type
        costs = summary.get("koszty_operacyjne")
        if student_count is not None and student_count != 0 and costs is not None:
            summary["koszt_na_ucznia"] = costs / student_count
        else:
            summary["koszt_na_ucznia"] = None

        report_rows.append({"placowka": facility_name, **summary})
        per_facility_tables[facility_name] = pd.DataFrame(rows)
        issues[facility_name] = detect_issues(facility_name, summary, student_count)

    # DataFrame zbiorczy
    summary_df = pd.DataFrame(report_rows)
    # Kolejność kolumn dla czytelności
    ordered_cols = [
        "placowka",
        "typ",
        "przychody_netto",
        "dotacje_podstawowe",
        "przychody_budzetowe",
        "koszty_operacyjne",
        "amortyzacja",
        "materialy_i_energia",
        "uslugi_obce",
        "podatki_i_oplaty",
        "wynagrodzenia",
        "ubezpieczenia_i_swiadczenia",
        "pozostale_koszty_rodzajowe",
        "pozostale_przychody_operacyjne",
        "pozostale_koszty_operacyjne",
        "zysk_strata_netto",
        "liczba_uczniow",
        "koszt_na_ucznia",
    ]
    summary_df = summary_df[ordered_cols]
    summary_df.sort_values("placowka", inplace=True)

    # Eksport do Excela: arkusz zbiorczy + arkusze placówek
    with pd.ExcelWriter(SUMMARY_XLSX, engine="openpyxl") as writer:
        summary_df.to_excel(writer, sheet_name="Zbiorcze_porownanie", index=False)
        for name, df in per_facility_tables.items():
            # skracamy nazwę arkusza do 31 znaków
            sheet_name = name[:31]
            df.to_excel(writer, sheet_name=sheet_name, index=False)

    # Dokument Word z uwagami
    doc = Document()
    doc.add_heading("Uwagi i potencjalne nieprawidłowości – sprawozdania 2024", level=1)
    for name in sorted(issues.keys()):
        doc.add_heading(name, level=2)
        for item in issues[name]:
            doc.add_paragraph(item, style="List Bullet")
    doc.save(ISSUES_DOCX)

    print(f"Zapisano raport Excel: {SUMMARY_XLSX}")
    print(f"Zapisano dokument Word: {ISSUES_DOCX}")


if __name__ == "__main__":
    main()
